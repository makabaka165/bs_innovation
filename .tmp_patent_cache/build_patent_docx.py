from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.shared import Cm, Inches, Pt


ROOT = Path(r"E:\bs_innovation")
TEMPLATE = ROOT / "zhuanli" / "一种二维平面阵列差分共阵孔径扩展方法-初稿(1).docx"
MARKDOWN = ROOT / "zhuanli" / "一种基于圆柱阵旋转等价性的规范波束域流形缓存方法、装置、设备、介质和程序.md"
OUTPUT = ROOT / "zhuanli" / "一种基于圆柱阵旋转等价性的规范波束域流形缓存方法、装置、设备、介质和程序.docx"

CN_FONT = "方正楷体_GB2312"
WEST_FONT = "Times New Roman"
BODY_SIZE = Pt(12)


def set_run_font(run, size=BODY_SIZE, bold=False, italic=False, cn_font=CN_FONT, west_font=WEST_FONT):
    run.font.name = west_font
    run.font.size = size
    run.font.bold = bold
    run.font.italic = italic
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), west_font)
    rfonts.set(qn("w:hAnsi"), west_font)
    rfonts.set(qn("w:eastAsia"), cn_font)
    rfonts.set(qn("w:cs"), west_font)


def set_paragraph_keep(paragraph, keep_next=False, keep_lines=True):
    ppr = paragraph._p.get_or_add_pPr()
    if keep_next:
        ppr.append(OxmlElement("w:keepNext"))
    if keep_lines:
        ppr.append(OxmlElement("w:keepLines"))
    ppr.append(OxmlElement("w:widowControl"))


def set_bottom_border(paragraph, size="10", space="6"):
    ppr = paragraph._p.get_or_add_pPr()
    pbdr = ppr.find(qn("w:pBdr"))
    if pbdr is None:
        pbdr = OxmlElement("w:pBdr")
        ppr.append(pbdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), space)
    bottom.set(qn("w:color"), "000000")
    pbdr.append(bottom)


def set_cell_margins(cell, top=60, start=80, bottom=60, end=80):
    tc = cell._tc
    tcpr = tc.get_or_add_tcPr()
    tcmar = tcpr.first_child_found_in("w:tcMar")
    if tcmar is None:
        tcmar = OxmlElement("w:tcMar")
        tcpr.append(tcmar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tcmar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tcmar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_shading(cell, fill):
    tcpr = cell._tc.get_or_add_tcPr()
    shd = tcpr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcpr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_repeat_table_header(row):
    trpr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    trpr.append(header)


def set_row_cant_split(row):
    trpr = row._tr.get_or_add_trPr()
    cant = OxmlElement("w:cantSplit")
    trpr.append(cant)


def set_table_borders_none(table):
    tblpr = table._tbl.tblPr
    borders = tblpr.find(qn("w:tblBorders"))
    if borders is not None:
        tblpr.remove(borders)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "nil")
        borders.append(element)
    tblpr.append(borders)


def set_table_borders_grid(table, size="6"):
    tblpr = table._tbl.tblPr
    borders = tblpr.find(qn("w:tblBorders"))
    if borders is not None:
        tblpr.remove(borders)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "000000")
        borders.append(element)
    tblpr.append(borders)


def add_page_field(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for item in (begin, instr, separate, text, end):
        run._r.append(item)
    set_run_font(run, size=Pt(10.5))


def restart_page_number(section):
    sectpr = section._sectPr
    for old in sectpr.findall(qn("w:pgNumType")):
        sectpr.remove(old)
    pg = OxmlElement("w:pgNumType")
    pg.set(qn("w:start"), "1")
    sectpr.append(pg)


def set_line_numbers(section):
    sectpr = section._sectPr
    for old in sectpr.findall(qn("w:lnNumType")):
        sectpr.remove(old)
    ln = OxmlElement("w:lnNumType")
    ln.set(qn("w:countBy"), "5")
    ln.set(qn("w:distance"), "360")
    ln.set(qn("w:restart"), "newPage")
    sectpr.append(ln)


def clear_story(story):
    element = story._element
    for child in list(element):
        element.remove(child)
    story.add_paragraph()


def clear_paragraph_borders(paragraph):
    ppr = paragraph._p.get_or_add_pPr()
    pbdr = ppr.find(qn("w:pBdr"))
    if pbdr is not None:
        ppr.remove(pbdr)
    pstyle = ppr.find(qn("w:pStyle"))
    if pstyle is not None:
        ppr.remove(pstyle)


def configure_section(section):
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.orientation = 0
    section.left_margin = Inches(0.98)
    section.right_margin = Inches(0.79)
    section.top_margin = Inches(0.98)
    section.bottom_margin = Inches(0.59)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.28)
    section.gutter = 0
    section.different_first_page_header_footer = False

    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    clear_story(section.header)
    clear_story(section.footer)
    header_p = section.header.paragraphs[0]
    header_p.clear()
    clear_paragraph_borders(header_p)
    header_p.paragraph_format.space_before = Pt(0)
    header_p.paragraph_format.space_after = Pt(0)
    footer_p = section.footer.paragraphs[0]
    footer_p.clear()
    add_page_field(footer_p)
    restart_page_number(section)
    set_line_numbers(section)


def clean_template_document():
    doc = Document(TEMPLATE)
    body = doc._element.body
    sectpr = body.sectPr
    for child in list(body):
        if child is not sectpr:
            body.remove(child)

    for rid, rel in list(doc.part.rels.items()):
        if rel.reltype == RT.IMAGE:
            del doc.part.rels[rid]

    normal = doc.styles["Normal"]
    normal.font.name = WEST_FONT
    normal.font.size = BODY_SIZE
    rpr = normal.element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), WEST_FONT)
    rfonts.set(qn("w:hAnsi"), WEST_FONT)
    rfonts.set(qn("w:eastAsia"), CN_FONT)
    rfonts.set(qn("w:cs"), WEST_FONT)

    for style in doc.styles:
        if style.type == 1 and style.name.lower() == "header":
            ppr = style.element.get_or_add_pPr()
            pbdr = ppr.find(qn("w:pBdr"))
            if pbdr is not None:
                ppr.remove(pbdr)

    doc.settings.odd_and_even_pages_header_footer = False
    doc.core_properties.title = "一种基于圆柱阵旋转等价性的规范波束域流形缓存方法、装置、设备、介质和程序"
    doc.core_properties.subject = "规范波束域流形缓存专利申请文件"
    doc.core_properties.keywords = "圆柱阵; 波束域; 规范缓存; 旋转等价; 精确网格查表"
    doc.core_properties.comments = "公式保留为LaTeX文本，便于后续MathType转换。"
    configure_section(doc.sections[0])
    return doc


def style_body_paragraph(paragraph, first_indent=True, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    paragraph.alignment = align
    fmt = paragraph.paragraph_format
    fmt.left_indent = Pt(0)
    fmt.right_indent = Pt(0)
    fmt.first_line_indent = Pt(24) if first_indent else Pt(0)
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    fmt.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    fmt.line_spacing = Pt(18)
    set_paragraph_keep(paragraph, keep_lines=False)


def add_runs_with_font(paragraph, text, size=BODY_SIZE, bold=False, italic=False, cn_font=CN_FONT, west_font=WEST_FONT):
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold, italic=italic, cn_font=cn_font, west_font=west_font)
    return run


def add_major_title(doc, text):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(12)
    fmt.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    fmt.line_spacing = Pt(28)
    add_runs_with_font(paragraph, text, size=Pt(18))
    set_bottom_border(paragraph)
    set_paragraph_keep(paragraph, keep_next=True)
    return paragraph


def add_invention_title(doc, text):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(6)
    fmt.space_after = Pt(10)
    fmt.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    fmt.line_spacing = Pt(24)
    add_runs_with_font(paragraph, text, size=Pt(15), bold=True)
    set_paragraph_keep(paragraph, keep_next=True)
    return paragraph


def add_section_heading(doc, text):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(9)
    fmt.space_after = Pt(6)
    fmt.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    fmt.line_spacing = Pt(20)
    add_runs_with_font(paragraph, text, size=Pt(12), bold=True)
    set_paragraph_keep(paragraph, keep_next=True)
    return paragraph


def add_subheading(doc, text, level):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    fmt = paragraph.paragraph_format
    fmt.first_line_indent = Pt(0)
    fmt.space_before = Pt(8 if level == 3 else 5)
    fmt.space_after = Pt(3)
    fmt.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    fmt.line_spacing = Pt(18)
    add_runs_with_font(paragraph, text, size=Pt(12), bold=True)
    set_paragraph_keep(paragraph, keep_next=True)
    return paragraph


def add_body(doc, text, current_major):
    paragraph = doc.add_paragraph()
    stripped = text.strip()
    is_claim_head = bool(re.match(r"^\d+、", stripped))
    is_caption = bool(re.fullmatch(r"图\s*\d+", stripped))
    if is_caption:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fmt = paragraph.paragraph_format
        fmt.first_line_indent = Pt(0)
        fmt.space_before = Pt(4)
        fmt.space_after = Pt(4)
        fmt.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        fmt.line_spacing = Pt(18)
        add_runs_with_font(paragraph, stripped, size=Pt(11))
        set_paragraph_keep(paragraph, keep_lines=True)
        return paragraph

    first_indent = not is_claim_head
    style_body_paragraph(paragraph, first_indent=first_indent)
    add_runs_with_font(paragraph, stripped)
    if is_claim_head:
        paragraph.paragraph_format.space_before = Pt(5)
        set_paragraph_keep(paragraph, keep_next=True)
    elif current_major == "权利要求书":
        paragraph.paragraph_format.left_indent = Pt(0)
        paragraph.paragraph_format.first_line_indent = Pt(24)
    return paragraph


def add_equation(doc, formula_lines):
    formula = " ".join(line.strip() for line in formula_lines if line.strip())
    tag_match = re.search(r"\\tag\{([^}]+)\}", formula)
    tag = tag_match.group(1) if tag_match else None
    if tag_match:
        formula = (formula[: tag_match.start()] + formula[tag_match.end() :]).strip()

    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders_none(table)
    widths = [Inches(0.42), Inches(5.45), Inches(0.42)]
    for idx, width in enumerate(widths):
        table.columns[idx].width = width
        table.cell(0, idx).width = width
        set_cell_margins(table.cell(0, idx), top=15, start=15, bottom=15, end=15)

    center = table.cell(0, 1)
    center.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    cp = center.paragraphs[0]
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp.paragraph_format.first_line_indent = Pt(0)
    cp.paragraph_format.space_before = Pt(2)
    cp.paragraph_format.space_after = Pt(2)
    cp.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    add_runs_with_font(cp, formula, size=Pt(10.5))

    if tag:
        rp = table.cell(0, 2).paragraphs[0]
        rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        rp.paragraph_format.first_line_indent = Pt(0)
        add_runs_with_font(rp, f"（{tag}）", size=Pt(10.5))
    set_row_cant_split(table.rows[0])
    return table


def add_markdown_table(doc, rows):
    parsed = []
    for line in rows:
        parsed.append([cell.strip() for cell in line.strip().strip("|").split("|")])
    if len(parsed) >= 2 and all(re.fullmatch(r":?-{3,}:?", cell.replace(" ", "")) for cell in parsed[1]):
        parsed.pop(1)
    column_count = max(len(row) for row in parsed)
    table = doc.add_table(rows=len(parsed), cols=column_count)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders_grid(table)
    widths = [Inches(2.65), Inches(3.6)] if column_count == 2 else [Inches(6.25 / column_count)] * column_count
    for r_idx, source_row in enumerate(parsed):
        row = table.rows[r_idx]
        set_row_cant_split(row)
        for c_idx in range(column_count):
            cell = row.cells[c_idx]
            cell.width = widths[min(c_idx, len(widths) - 1)]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell, top=65, start=80, bottom=65, end=80)
            if r_idx == 0:
                set_cell_shading(cell, "E7E6E6")
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if r_idx == 0 else WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.first_line_indent = Pt(0)
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            paragraph.paragraph_format.line_spacing = Pt(15)
            value = source_row[c_idx] if c_idx < len(source_row) else ""
            add_runs_with_font(paragraph, value, size=Pt(10.5), bold=(r_idx == 0))
    if table.rows:
        set_repeat_table_header(table.rows[0])
    return table


def add_code_block(doc, lines):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders_grid(table)
    cell = table.cell(0, 0)
    set_cell_margins(cell, top=80, start=120, bottom=80, end=120)
    cell.text = ""
    for idx, line in enumerate(lines):
        paragraph = cell.paragraphs[0] if idx == 0 else cell.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.first_line_indent = Pt(0)
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        paragraph.paragraph_format.line_spacing = Pt(13)
        add_runs_with_font(paragraph, line, size=Pt(9), cn_font="楷体", west_font="Courier New")
    set_row_cant_split(table.rows[0])
    return table


def set_picture_alt_text(inline_shape, alt_text):
    docpr = inline_shape._inline.docPr
    docpr.set("descr", alt_text)
    docpr.set("title", alt_text)


def add_image(doc, image_path, alt_text, current_major, figure_index):
    if current_major == "说明书附图" and figure_index > 0:
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.add_run().add_break(WD_BREAK.PAGE)

    width = Inches(5.15)
    if "图1_" not in image_path.name:
        width = Inches(6.3)
    elif current_major == "说明书附图":
        width = Inches(5.25)
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Pt(0)
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.space_after = Pt(2)
    run = paragraph.add_run()
    shape = run.add_picture(str(image_path), width=width)
    set_picture_alt_text(shape, alt_text)
    set_paragraph_keep(paragraph, keep_lines=True)
    return paragraph


def normalize_major(text):
    return text.replace("　", "").replace(" ", "")


def build():
    doc = clean_template_document()
    lines = MARKDOWN.read_text(encoding="utf-8").splitlines()
    current_major = ""
    drawing_figure_index = 0
    i = 0
    while i < len(lines):
        raw = lines[i]
        line = raw.strip()
        if not line:
            i += 1
            continue
        if line == "<!-- WORD_SECTION -->":
            section = doc.add_section(WD_SECTION.NEW_PAGE)
            configure_section(section)
            i += 1
            continue
        if line == "<!-- WORD_PAGE_BREAK -->":
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.add_run().add_break(WD_BREAK.PAGE)
            i += 1
            continue
        if line.startswith("# "):
            title = line[2:].strip()
            current_major = normalize_major(title)
            drawing_figure_index = 0
            add_major_title(doc, title)
            i += 1
            continue
        if line.startswith("## "):
            title = line[3:].strip()
            if title.startswith("一种基于圆柱阵"):
                add_invention_title(doc, title)
            else:
                add_section_heading(doc, title)
            i += 1
            continue
        if line.startswith("### "):
            add_subheading(doc, line[4:].strip(), 3)
            i += 1
            continue
        if line.startswith("#### "):
            add_subheading(doc, line[5:].strip(), 4)
            i += 1
            continue
        if line == "$$":
            formula_lines = []
            i += 1
            while i < len(lines) and lines[i].strip() != "$$":
                formula_lines.append(lines[i])
                i += 1
            add_equation(doc, formula_lines)
            i += 1
            continue
        if line.startswith("~~~"):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("~~~"):
                code_lines.append(lines[i])
                i += 1
            add_code_block(doc, code_lines)
            i += 1
            continue
        image_match = re.fullmatch(r"!\[([^\]]*)\]\(([^)]+)\)", line)
        if image_match:
            alt_text = image_match.group(1)
            image_path = (MARKDOWN.parent / image_match.group(2)).resolve()
            add_image(doc, image_path, alt_text, current_major, drawing_figure_index)
            if current_major == "说明书附图":
                drawing_figure_index += 1
            i += 1
            continue
        if line.startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1
            add_markdown_table(doc, table_lines)
            continue

        add_body(doc, raw, current_major)
        i += 1

    for section in doc.sections:
        configure_section(section)
    doc.save(OUTPUT)
    print(OUTPUT)
    print(f"sections={len(doc.sections)}")
    print(f"paragraphs={len(doc.paragraphs)}")
    print(f"tables={len(doc.tables)}")
    print(f"inline_shapes={len(doc.inline_shapes)}")


if __name__ == "__main__":
    build()
